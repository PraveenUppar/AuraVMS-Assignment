import pytest
from app.parser.txt_md import parse_txt_md
from app.parser.base import ParseError
from app.parser.docx import parse_docx



def test_simple_text():
    text = "My Title\nThis is content."
    result = parse_txt_md(text)
    assert result.title == "My Title"
    assert result.content == "This is content."
    assert result.image is None


def test_image_first_line():
    text = "![alt](image.png)\nMy Title\nBody"
    result = parse_txt_md(text)
    assert result.image == "![alt](image.png)"
    assert result.title == "My Title"


def test_empty_file():
    with pytest.raises(ParseError):
        parse_txt_md("   ")

def test_docx_parsing(tmp_path):
    from docx import Document

    docx_file = tmp_path / "sample.docx"

    doc = Document()
    doc.add_paragraph("My DOCX Title")
    doc.add_paragraph("DOCX content here")
    doc.save(docx_file)

    result = parse_docx(str(docx_file))

    assert result.title == "My DOCX Title"
    assert result.content == "DOCX content here"